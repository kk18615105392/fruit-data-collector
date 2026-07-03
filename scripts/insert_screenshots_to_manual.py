# -*- coding: utf-8 -*-
"""将 screenshots 目录中的截图插入用户手册 Word 文档（自动缩放）"""

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "data" / "软著材料"
SCREEN_DIR = OUT_DIR / "screenshots"
MANUAL_PATH = OUT_DIR / "热带水果数据集采集系统-用户手册.docx"

# Word 中图片最大宽度/高度（厘米）
MAX_WIDTH_CM = 9.0
MAX_HEIGHT_CM = 13.5

# 图号 -> 截图文件；可按图单独微调宽度
FIGURE_MAP = {
    "图4-2": ("fig4-2-home.png", 8.5),
    "图4-3": ("fig4-3-collect.png", 8.5),
    "图4-4": ("fig4-4-list.png", 8.5),
    "图4-5": ("fig4-5-detail.png", 8.0),
    "图4-6": ("fig4-6-export.png", 8.5),
    "图5-1": ("fig4-2-home.png", 8.5),
    "图5-2": ("fig4-3-collect.png", 8.5),
    "图5-3": ("fig4-4-list.png", 8.5),
    "图5-4": ("fig4-6-export.png", 8.5),
    "图5-5": ("fig5-5-disease.png", 8.5),
}


def calc_picture_size(img_path: Path, target_width_cm: float) -> tuple[float, float]:
    with Image.open(img_path) as im:
        px_w, px_h = im.size
    if px_w <= 0:
        return target_width_cm, target_width_cm * 0.75

    aspect = px_h / px_w
    width_cm = target_width_cm
    height_cm = width_cm * aspect
    if height_cm > MAX_HEIGHT_CM:
        height_cm = MAX_HEIGHT_CM
        width_cm = height_cm / aspect
    return width_cm, height_cm


def compress_for_word(img_path: Path, max_px: int = 900) -> BytesIO:
    """压缩像素尺寸，减小 docx 体积"""
    with Image.open(img_path) as im:
        im = im.convert("RGB") if im.mode in ("RGBA", "P") else im
        w, h = im.size
        if w > max_px:
            h = int(h * max_px / w)
            w = max_px
            im = im.resize((w, h), Image.Resampling.LANCZOS)
        buf = BytesIO()
        im.save(buf, format="JPEG", quality=82, optimize=True)
        buf.seek(0)
        return buf


def insert_figures() -> None:
    if not MANUAL_PATH.exists():
        raise FileNotFoundError(f"未找到用户手册: {MANUAL_PATH}")

    doc = Document(str(MANUAL_PATH))
    inserted: set[str] = set()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text.startswith("图"):
            continue

        fig_key = text.split()[0]
        if fig_key not in FIGURE_MAP or fig_key in inserted:
            continue

        img_name, width_cm = FIGURE_MAP[fig_key]
        img_path = SCREEN_DIR / img_name
        if not img_path.exists():
            print(f"跳过 {fig_key}，缺少 {img_name}")
            continue

        w_cm, h_cm = calc_picture_size(img_path, width_cm)
        img_buf = compress_for_word(img_path)

        parent = para._element.getparent()
        idx = list(parent).index(para._element)

        new_p = doc.add_paragraph()
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = new_p.add_run()
        run.add_picture(img_buf, width=Cm(w_cm), height=Cm(h_cm))
        parent.insert(idx + 1, new_p._element)

        inserted.add(fig_key)
        print(f"已插入 {fig_key} <- {img_name} ({w_cm:.1f}x{h_cm:.1f} cm)")

    out_path = OUT_DIR / "热带水果数据集采集系统-用户手册-含截图.docx"
    try:
        doc.save(str(out_path))
    except PermissionError:
        out_path = OUT_DIR / "热带水果数据集采集系统-用户手册-含截图-更新.docx"
        doc.save(str(out_path))
        print("原文件被占用，已另存为新文件。")
    print(f"\n已保存: {out_path}")


if __name__ == "__main__":
    insert_figures()
