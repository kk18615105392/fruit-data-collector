# -*- coding: utf-8 -*-
"""生成软著三件套：源代码、用户手册、软件设计说明书"""

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "data" / "软著材料"
SYSTEM_NAME = "热带水果采集APP"
VERSION = "V1.0"
FULL_NAME = f"{SYSTEM_NAME}{VERSION}"

SOURCE_FILES = [
    "src/App.tsx",
    "src/main.tsx",
    "src/types.ts",
    "src/constants.ts",
    "src/disease.ts",
    "src/storage.ts",
    "src/naming.ts",
    "src/namingSettings.ts",
    "src/photoUtils.ts",
    "src/fileStorage.ts",
    "src/fsUtils.ts",
    "src/exportGroups.ts",
    "src/export.ts",
    "src/components/CollectForm.tsx",
    "src/components/HomePage.tsx",
    "src/components/RecordList.tsx",
    "src/components/RecordDetail.tsx",
    "src/components/ExportPage.tsx",
    "src/components/BottomNav.tsx",
    "capacitor.config.ts",
    "android/app/src/main/java/com/fruitdata/collector/MainActivity.java",
]


def set_doc_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_body(doc: Document, text: str) -> None:
    for line in text.strip().split("\n"):
        p = doc.add_paragraph(line.strip())
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5


def _style_table_cell(cell, *, bold: bool = False, align_center: bool = True) -> None:
    for paragraph in cell.paragraphs:
        if align_center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = bold
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(12)


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(12)


def save_doc(doc: Document, path: Path) -> Path:
    try:
        doc.save(str(path))
        return path
    except PermissionError:
        alt = path.with_stem(f"{path.stem}-更新")
        doc.save(str(alt))
        print(f"原文件被占用，已另存: {alt.name}")
        return alt


def add_data_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        _style_table_cell(cell, bold=True)

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = value
            _style_table_cell(cell)

    doc.add_paragraph()


def generate_source_code_doc() -> Path:
    doc = Document()
    set_doc_style(doc)
    add_title(doc, f"{FULL_NAME}源代码")

    for rel in SOURCE_FILES:
        path = ROOT / rel
        if not path.exists():
            continue
        content = path.read_text(encoding="utf-8")
        doc.add_paragraph(f"// ===== {rel} =====")
        doc.add_paragraph(content)

    out = OUT_DIR / f"{SYSTEM_NAME}-源代码.docx"
    return save_doc(doc, out)


def generate_user_manual() -> Path:
    doc = Document()
    set_doc_style(doc)
    add_title(doc, f"{FULL_NAME}用户手册")

    add_heading(doc, "1. 系统简介", 1)
    add_heading(doc, "1.1 编写目的", 2)
    add_body(
        doc,
        """
本手册是《热带水果采集APP V1.0》的配套使用说明，旨在帮助农业科研人员、数据集标注人员、智慧农业项目工作人员及相关从业者全面了解软件的功能定位、核心技术、操作流程与实际应用场景。通过本手册，用户可快速掌握如何利用移动端应用实现热带水果图像样本的连续采集、病害标注、多维属性记录、按规则命名保存及数据集导出，解决传统人工记录分散、命名不规范、病害标签缺失、图片与标签难以统一管理的问题，为机器学习模型训练、病害识别研究及种质资源调查提供标准化数据支撑。
        """,
    )

    add_heading(doc, "1.2 项目背景", 2)
    add_body(
        doc,
        """
在热带水果种质资源研究、病害识别模型训练及智慧农业示范项目中，高质量标注数据集是算法研发的基础。传统数据采集方式依赖相机拍照后再用 Excel 或纸质表格记录，存在以下问题：一是图片与标签分离，后期整理成本高；二是文件命名随意，难以批量导入训练框架；三是现场采集效率低，无法连续快速录入多个样本；四是缺乏病害类型、GPS、成熟度等结构化字段，数据价值受限；五是现场人员难以准确判断病害时缺少「未知」等标准化标注选项。

为解决上述痛点，本软件面向 Android 移动终端，构建了一套集「连续拍照—病害标注—属性记录—规则命名—本地存储—批量导出」于一体的热带水果采集应用。软件内置芒果、菠萝、火龙果、榴莲、山竹等 18 种常见热带水果类别，每种水果均配置常见病害选项；支持选择「未知」表示不确定病害类型，支持「上一个病害」快速沿用最近标签。软件同时支持颜色、成熟度、重量、备注及定位信息采集，并可将样本直接保存至手机 FruitCollector 相册，同时生成 JSON、CSV 格式数据集，适用于 YOLO、TensorFlow 等深度学习框架的数据预处理流程。
        """,
    )

    add_heading(doc, "2. 系统概述", 1)
    add_heading(doc, "2.1 设计目标", 2)
    add_body(
        doc,
        """
① 构建轻量化、易操作的热带水果移动端采集平台；
② 支持连续拍照与批次属性统一标注，提升现场采集效率；
③ 按水果种类提供常见病害选项，支持「未知」与「上一个病害」快捷标注；
④ 实现基于用户所选属性的自动文件命名与自定义命名；
⑤ 将图片及元数据直接保存至手机存储，便于离线使用与备份；
⑥ 提供数据浏览、搜索、编辑、删除及一键导出能力。
        """,
    )

    add_heading(doc, "2.2 功能特点", 2)
    add_body(
        doc,
        """
热带水果采集APP是一套专注于热带水果图像样本采集与管理的移动应用，融合 React 跨平台前端与 Capacitor 原生能力，具备以下核心功能：

连续采集模式：用户先设置本批次样本属性（种类、病害、颜色、成熟度等），再通过「继续拍照」按钮连续拍摄多张水果照片，无需重复填写公共字段，适合批量现场采样；

病害智能标注：选择水果种类后，系统自动展示该水果的常见病害列表（如芒果的炭疽病、白粉病、疮痂病等）；用户可选择具体病害，也可选择「未知（不清楚具体病害）」或「上一个病害」快速沿用最近标签，以及「健康」表示无病害；

智能文件命名：系统支持按种类、病害、样本名、颜色、成熟度、重量、序号、时间戳等字段组合自动生成文件名，例如「芒果_炭疽病_黄色_成熟_001.jpg」，并支持名称前缀、后缀及自定义文件名覆盖；

手机本地存储：保存时优先写入手机相册「FruitCollector」，同时在应用私有目录保留原图副本及 batch_xxx.json 批次元数据，可在图库或应用数据中查看；

多维样本标注：每条样本包含水果种类、病害类型、重量、颜色、成熟度、GPS 坐标、备注及时间戳，满足病害识别与品质分析模型训练对标签丰富度的要求；

数据管理与导出：应用内提供列表浏览、分类筛选、关键词搜索、详情查看与编辑删除；导出模块支持自定义数据集名称、按采集批次勾选导出范围，并生成 JSON + CSV + 图片文件夹，CSV 含 disease 字段，便于后续模型训练与统计分析。
        """,
    )

    add_heading(doc, "3. 运行环境", 1)
    add_heading(doc, "3.1 硬件运行环境", 2)
    add_table_caption(doc, "表1 硬件运行环境")
    add_data_table(
        doc,
        ["选项", "要求"],
        [
            ["处理器", "ARM 64 位及以上移动处理器"],
            ["内存", "2GB 及以上"],
            ["存储空间", "100MB 应用安装空间；采集数据另计"],
            ["摄像头", "支持后置/前置相机"],
            ["定位", "支持 GPS 或网络定位（可选）"],
            ["网络", "导出与云端构建需网络；现场采集可离线"],
        ],
    )

    add_heading(doc, "3.2 系统运行环境", 2)
    add_table_caption(doc, "表2 系统运行环境")
    add_data_table(
        doc,
        ["用途", "要求"],
        [
            ["操作系统", "Android 6.0（API 23）及以上"],
            ["应用形式", "Android APK 安装包"],
            ["开发框架", "React 18 + TypeScript + Vite"],
            ["原生容器", "Capacitor 7"],
            ["本地存储", "LocalStorage + Android 文件系统"],
        ],
    )

    add_heading(doc, "4. 系统架构设计", 1)
    add_heading(doc, "4.1 功能模块架构", 2)
    add_body(
        doc,
        """
热带水果采集APP基于模块化设计思想构建，主要划分为以下核心功能模块：

图4-1 热带水果采集APP功能模块

（1）首页统计模块：展示已采集样本总数、水果种类数、含图片样本数、含定位样本数，并提供快速入口跳转至采集、列表与导出页面。

（2）连续采集模块：采集流程分为三步——①设置批次属性（种类、病害类型、重量、颜色、成熟度、备注、GPS）；②配置命名规则（种类、病害、颜色等字段勾选，前缀后缀、自定义文件名、实时预览）；③连续拍照并网格预览，支持删除单张与一键保存到手机。

（3）病害标注模块：内置 FRUIT_DISEASES 病害字典，按 18 种热带水果分别配置常见病害；提供「未知」「上一个病害」「健康」及具体病害名称等选项；resolveDisease 函数解析用户选择并写入 FruitRecord.disease 字段，同时更新 lastDisease 缓存供下次快速沿用。

（4）数据管理模块：以卡片列表形式展示历史样本，支持按种类筛选与关键词搜索，进入详情页可查看文件名、病害、手机保存路径、批次 ID、图片及全部标注字段，并支持编辑与删除。

（5）文件存储模块：通过 Capacitor Media 与 Filesystem 接口将图片优先保存至 FruitCollector 相册，并在 Directory.Data 应用私有目录保留副本；fsUtils 模块处理目录已存在等边界情况，按命名规则生成独立图片文件及批次 JSON 元数据。

（6）数据集导出模块：ExportPage 组件提供导出名称输入框与批次选择列表，exportGroups 模块按 batchId 将样本分组为可勾选的数据集；用户可自定义导出名称，勾选需导出的采集批次后，系统将选中样本导出为 JSON 元数据、CSV 表格及 images 图片目录，Android 端可通过系统分享功能发送数据集包。
        """,
    )

    add_heading(doc, "4.2 界面设计", 2)
    add_heading(doc, "4.2.1 首页界面", 3)
    add_body(
        doc,
        """
首页顶部为绿色渐变横幅，标题为「热带水果采集APP」，副标题说明「连续拍照 · 属性命名 · 一键保存到手机相册目录」。中部为四宫格统计卡片，分别显示已采集样本数、水果种类数、含图片样本数、含定位样本数。下方卡片展示支持采集的热带水果标签及三个快捷按钮：「开始采集新样本」「查看已采集数据」「导出数据集」。底部为固定导航栏，包含首页、采集、数据、导出四个 Tab。

图4-2 热带水果采集APP首页
        """,
    )

    add_heading(doc, "4.2.2 连续采集界面", 3)
    add_body(
        doc,
        """
连续采集页分为三个区域。区域一「样本属性」包含水果种类（必填）、病害类型（随种类动态加载常见病害，含「未知」「上一个病害」「健康」选项）、样本名称、重量、颜色、成熟度、备注及 GPS 定位按钮。区域二「文件命名规则」可展开配置，提供种类、病害、样本名、颜色、成熟度、重量、序号、时间戳等命名字段勾选，以及前缀、后缀、自定义文件名输入框和「下一张预览」实时显示（如「芒果_炭疽病_黄色_成熟_001.jpg」）。区域三「连续拍照」显示已拍张数，提供橙色「继续拍照」主按钮、相册添加按钮、照片网格预览（含序号与文件名）及「保存 N 张到手机」提交按钮。

图4-3 连续采集界面
        """,
    )

    add_heading(doc, "4.2.3 数据列表与详情界面", 3)
    add_body(
        doc,
        """
数据列表页顶部显示记录总数，提供搜索框与种类筛选 Chip。每条记录以缩略图、名称、文件名/病害/成熟度摘要展示，点击进入详情页。详情页大图展示样本图片，下方列表显示文件名、手机路径、样本 ID、重量、颜色、病害、成熟度、GPS 坐标、采集时间与备注，底部提供编辑与删除按钮。

图4-4 数据列表界面
图4-5 样本详情界面
        """,
    )

    add_heading(doc, "4.2.4 导出界面", 3)
    add_body(
        doc,
        """
导出页分为三个区域。区域一「导出名称」提供数据集名称输入框，用户可自定义如「芒果病害数据集_20260703」，该名称将作为导出文件夹及 JSON 文件名。区域二「选择数据集」按采集批次（batchId）列出历史样本组，每组显示缩略图、水果名称、病害标签、样本数量与采集时间，支持复选框单选/多选及「全选/取消全选」。区域三「导出内容」展示已选批次数、样本总数及 JSON/CSV/图片说明，点击「导出 N 条样本」按钮后，Android 端生成完整数据包并通过系统分享界面发送。

图4-6 数据集导出界面
        """,
    )

    add_heading(doc, "5. 操作示例", 1)
    add_heading(doc, "5.1 安装与启动", 2)
    add_body(
        doc,
        """
用户获取 app-debug.apk 安装包后，在 Android 手机设置中允许「未知来源」安装，完成安装。首次启动时，系统请求相机、存储及定位权限，用户点击「允许」后即可进入首页。

图5-1 应用安装与权限授权
        """,
    )

    add_heading(doc, "5.2 连续采集操作", 2)
    add_body(
        doc,
        """
步骤一：点击底部导航「采集」进入连续采集页。步骤二：在「样本属性」中选择水果种类（如「芒果」），在「病害类型」中选择具体病害（如「炭疽病」）；若不清楚病害可选「未知」，若与上一张相同可选「上一个病害」。步骤三：可选填颜色「黄色」、成熟度「成熟」及备注。步骤四：展开「文件命名规则」，确认预览文件名为「芒果_炭疽病_黄色_成熟_001.jpg」。步骤五：点击「继续拍照」，对准水果拍摄，可连续拍摄多张，网格区实时显示缩略图与文件名。步骤六：点击「保存 3 张到手机」，系统将图片写入 FruitCollector 相册并同步至应用数据列表，弹出成功提示，可继续采集下一批。

图5-2 连续采集操作流程
        """,
    )

    add_heading(doc, "5.3 查看与管理数据", 2)
    add_body(
        doc,
        """
用户点击底部「数据」进入列表页，可通过搜索框输入「芒果」筛选，或点击种类 Chip 快速过滤。点击某条记录进入详情页查看完整标注与手机保存路径。如需修改，点击「编辑」返回采集表单；如需删除，点击「删除」并在确认框中确认。

图5-3 数据查看与管理
        """,
    )

    add_heading(doc, "5.4 导出数据集", 2)
    add_body(
        doc,
        """
用户点击底部「导出」，在导出页输入自定义数据集名称（如「芒果病害数据集_20260703」），在「选择数据集」区域勾选需导出的采集批次（可按批次只导出芒果或香蕉等子集），确认已选样本数量后点击「导出 N 条样本」。Android 端以用户输入的名称创建导出文件夹，生成包含 dataset.json、dataset.csv 及 images 子目录的数据集，并唤起系统分享面板，用户可选择发送至微信、邮件或保存至云盘，供后续模型训练使用。

图5-4 数据集导出操作
        """,
    )

    add_heading(doc, "5.5 病害标注说明", 2)
    add_table_caption(doc, "表3 病害标注选项说明")
    add_data_table(
        doc,
        ["选项", "含义", "适用场景"],
        [
            ["未知", "不清楚具体病害类型", "现场无法判断病害名称"],
            ["上一个病害", "沿用最近一次保存的病害标签", "连续拍摄同一病害样本"],
            ["健康", "确认无病害", "正常果实对照样本"],
            ["具体病害名", "如炭疽病、白粉病等", "能明确判断病害类型"],
        ],
    )
    add_body(
        doc,
        """
选择水果种类后，系统自动加载该水果的常见病害列表。例如芒果包括：炭疽病、白粉病、疮痂病、细菌性溃疡病、煤烟病、生理性病害、健康；香蕉包括：巴拿马病、叶斑病、黑星病、束顶病、健康等。

图5-5 病害类型选择界面
        """,
    )

    out = OUT_DIR / f"{SYSTEM_NAME}-用户手册.docx"
    return save_doc(doc, out)


def generate_design_doc() -> Path:
    doc = Document()
    set_doc_style(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("本材料用于软件著作权登记及相关成果认定。")

    doc.add_paragraph()
    add_title(doc, f"{FULL_NAME}软件开发说明")

    add_body(
        doc,
        """
课题来源：热带水果种质资源数字化采集与数据集建设
课题名称：热带水果采集APP研发
课题编号：（请填写）
        """,
    )

    add_body(doc, "课题负责人：（请填写）")
    add_body(doc, "课题起止时间：2025 年 01 月 01 日——2026 年 12 月 31 日")

    doc.add_paragraph()
    add_table_caption(doc, "表1 主要完成人及分工")
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rows = [
        ("姓名", "分工"),
        ("（请填写）", "软件总体设计、Android 端集成"),
        ("（请填写）", "前端界面与采集流程开发"),
        ("（请填写）", "数据存储、命名规则与导出模块开发"),
    ]
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        _style_table_cell(table.rows[i].cells[0], bold=(i == 0))
        _style_table_cell(table.rows[i].cells[1], bold=(i == 0))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("（请填写单位名称）")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("2026 年 07 月 03 日")

    add_heading(doc, "一、软件概述", 1)
    add_body(
        doc,
        """
热带水果采集APP V1.0 是一款运行于 Android 平台的移动应用软件，包名 com.fruitdata.collector，面向热带水果图像数据集的现场采集、标注、命名、存储与导出场景。软件采用 B/S 架构思想结合 Hybrid App 技术路线，使用 React + TypeScript 实现业务界面，通过 Capacitor 框架封装为原生 APK，调用 Android 相机、文件系统、定位及分享等系统能力。
        """,
    )

    add_heading(doc, "二、开发环境与工具", 1)
    add_table_caption(doc, "表1 开发环境与工具")
    add_data_table(
        doc,
        ["类别", "说明"],
        [
            ["开发语言", "TypeScript、JavaScript、Java、HTML、CSS"],
            ["前端框架", "React 18、Vite 5"],
            ["原生桥接", "Capacitor 7（Camera、Filesystem、Geolocation、Share、Media）"],
            ["Android 构建", "Gradle 8、Android SDK 35"],
            ["版本管理", "Git、GitHub Actions 云端构建"],
        ],
    )

    add_heading(doc, "三、系统结构设计", 1)
    add_body(
        doc,
        """
系统分为表现层、业务逻辑层与数据持久层。表现层包括 HomePage、CollectForm、RecordList、RecordDetail、ExportPage 等 React 组件；业务逻辑层包括 disease（病害字典与解析）、naming（命名规则引擎）、fileStorage（相册优先保存与本地副本）、fsUtils（文件系统安全写入）、photoUtils（缩略图压缩）、exportGroups（导出批次分组）、export（数据集导出）、storage（本地记录与 lastDisease 缓存）；数据持久层包括浏览器 LocalStorage 缩略图与元数据，以及 Android Directory.Data 原图副本。
        """,
    )

    add_heading(doc, "四、核心算法与处理流程", 1)
    add_body(
        doc,
        """
（1）命名生成算法：根据用户勾选的 NamingField 列表，依次提取种类、病害、颜色、成熟度、序号等字段值，经 sanitizeFileName 过滤非法字符后以指定分隔符拼接，并在批量模式下强制附加三位序号防止文件名冲突。

（2）病害解析算法：FRUIT_DISEASES 常量维护 18 种水果与常见病害的映射关系；resolveDisease 函数处理三种情况——选择具体病害名直接写入、选择 PREVIOUS_DISEASE_VALUE 时读取 lastDisease 缓存、选择 UNKNOWN_DISEASE（未知）时写入标准化「未知」标签。

（3）连续采集流程：SessionPhoto 数组暂存当前批次照片，每次拍照调用 Capacitor Camera API 获取 DataUrl，经 resolveFormForNaming 解析病害后调用 buildPhotoFileName 生成文件名；保存时批量调用 saveRecordsToPhone 写入文件系统并 addRecords 持久化至 LocalStorage，同时更新 lastDisease。

（4）导出流程：groupRecordsIntoDatasets 按 batchId 将历史样本分组为可选数据集；ExportPage 接收用户输入的 datasetName 与勾选的批次 ID，调用 exportDataset 导出；buildExportBundle 构建含 datasetName、disease 字段的 JSON 元数据结构，Native 端以自定义名称创建文件夹并写入 dataset.json、dataset.csv 及 images 子目录，Web 端则生成含 Base64 嵌入的单文件 JSON。
        """,
    )

    add_heading(doc, "五、测试与运行", 1)
    add_body(
        doc,
        """
软件在 Android 10 及以上机型及浏览器环境完成功能测试，覆盖连续拍照、病害标注（含未知/上一个病害）、属性命名、手机存储、列表检索、编辑删除及数据集导出等核心用例。项目内置 E2E 自动化测试（scripts/e2e-test.mjs），验证多图保存与自定义名称导出流程。应用安装包通过 GitHub Actions 自动构建，生成 app-debug.apk 供部署测试。
        """,
    )

    out = OUT_DIR / f"{SYSTEM_NAME}-软件设计说明书.docx"
    return save_doc(doc, out)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    files = [
        generate_source_code_doc(),
        generate_user_manual(),
        generate_design_doc(),
    ]
    for f in files:
        print(f"已生成: {f}")


if __name__ == "__main__":
    main()
